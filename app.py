# from fastapi import FastAPI, Form, Request
# from fastapi.responses import HTMLResponse
# from fastapi.templating import Jinja2Templates

# app = FastAPI()

# # Initialize Jinja2 templates
# templates = Jinja2Templates(directory="templates")

# # Create a route for the resume creation form
# @app.post("/create_resume/")
# async def create_resume(request: Request, full_name: str = Form(...), email: str = Form(...), experience: str = Form(...)):
#     context = {
#         "full_name": full_name,
#         "email": email,
#         "experience": experience
#     }
#     return templates.TemplateResponse("resume_template.html", {"request": request, "context": context})

# # Route to render the create resume form
# @app.get("/create_resume_form/")
# async def create_resume_form(request: Request):
#     return templates.TemplateResponse("create_resume_form.html", {"request": request})



from fastapi import FastAPI, Form, Request, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from docx.shared import Inches,Pt,RGBColor
import json
import base64
import os
from typing import List
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

app = FastAPI()

# Initialize Jinja2 templates
templates = Jinja2Templates(directory="templates")

# Create a route for the resume creation form
@app.post("/create_resume/")
async def create_resume(request: Request,
                        professional_summary: str = Form(...),
                        technical_skills: str = Form(...),
                        work_history: str = Form(...),  # Required field
                        education: str = Form(...),
                        full_name: str = Form(...),
                        company_logo: UploadFile = File(...),
                        designation: str = Form(...)):

    # Deserialize the JSON data from the work_history field
    try:
        work_history_data = json.loads(work_history)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail="Invalid work history JSON")

    # Set the value of 'client' to be the same as 'company' if it's missing
    for entry in work_history_data:
        if "client" not in entry:
            entry["client"] = entry["company"]

    # Inside the loop that processes work history
    for entry in work_history_data:
        entry["client"] = entry.get("client", entry["company"])
        entry["role"] = entry.get("role", "")
        entry["responsibilities"] = entry.get("responsibilities", [])

    # Process the professional_summary to create bullet points
    professional_summary_bullets = professional_summary.split('\n')
    technical_skills_bullets = technical_skills.split('\n')

    # Encode the company logo as base64
    company_logo_data = company_logo.file.read()
    company_logo_base64 = base64.b64encode(company_logo_data).decode("utf-8")

    context = {
        "professional_summary": professional_summary_bullets,
        "technical_skills": technical_skills_bullets,
        "work_history": work_history_data,
        "education": education,
        "full_name": full_name,
        "company_logo_base64": company_logo_base64,
        "designation": designation,
    }
    
    # Create a new Word document
    doc = Document()
    
    # Set the font of the document to Times New Roman
    doc.styles['Normal'].font.name = 'Times New Roman'
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.font_size = Pt(12)
    
    # Add a section with a header
    section = doc.sections[0]
    header = section.header

    # Add a table with 1 row and 2 columns for the logo and name+designation
    table = header.add_table(rows=1, cols=2,width=Inches(6))
    table.autofit = False
    table.columns[0].width = Inches(1.5)  # Adjust the width for the logo column
    
    
    # Add the company logo to the left cell of the table
    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center the logo vertically
    p = left_cell.paragraphs[0]
    run = p.add_run()
    run.add_picture(company_logo.file, width=Inches(2), height=Inches(1))

    # Add the name and designation to the right cell of the table
    right_cell = table.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center the text vertically
    p = right_cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right-align the text
    p.add_run(full_name).bold = True
    p.add_run("\n" + designation).italic = True
    p.space_after = Pt(12)  # Adjust the spacing between the name and designation
    
    # Add a paragraph with underscores to simulate a solid line
    line_paragraph = header.add_paragraph("_" * 75)
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # Add a horizontal line below the merged cell
    #doc.add_paragraph("_" * 103)

    # Add professional summary
    header1 = doc.add_heading("PROFESSIONAL SUMMARY", level=2)
    for run in header1.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        
    # Create a new paragraph after the heading
    paragraph = doc.add_paragraph()
    # Add a blank line to the new paragraph
    paragraph.paragraph_format.space_after = Pt(9)

    # Add bullet points with reduced spacing
    for point in professional_summary_bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("•").bold = True  # Add bullet point
        p.add_run(point)
        p.paragraph_format.line_spacing = Pt(9)  # Adjust the line spacing here

    # Add technical skills
    header2=doc.add_heading("Technical Skills", level=2)
    for run in header2.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Create a new paragraph after the heading
    paragraph = doc.add_paragraph()
    # Add a blank line to the new paragraph
    paragraph.paragraph_format.space_after = Pt(6)   
    
    for skill in technical_skills_bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("•").bold = True
        p.add_run(skill)
        p.paragraph_format.line_spacing = Pt(9)  # Adjust the line spacing here

    # Add Work History section
    header3=doc.add_heading("Work History", level=2)
    for run in header3.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # Add a paragraph for spacing after the "Work History" header
    spacing_paragraph = doc.add_paragraph()
    spacing_paragraph.paragraph_format.space_after = Pt(3)  
    # Create a table for Work History
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set the width of the columns
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)

    # Set table style to display borders
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the content

    # Add table headers
    table_headers = table.rows[0].cells
    table_headers[0].text = "Company"
    table_headers[1].text = "Duration"
    
    # Add a paragraph for spacing after the table
    spacing_paragraph = doc.add_paragraph()
    spacing_paragraph.paragraph_format.line_spacing = Pt(9)
    
    # Initialize a flag to check if it's the first entry
    first_entry = True

    # Add Work History entries from context.work_history
    for entry in context["work_history"]:
        row_cells = table.add_row().cells
        row_cells[0].text = entry['company']
        row_cells[1].text = entry['duration']
        
        # Add spacing between table entries
        if not first_entry:
            doc.add_paragraph()

        # Add custom format for each entry
        client_duration_paragraph = doc.add_paragraph()
        client_duration_paragraph.add_run(f"Client: {entry['company']}\t\t\t\t\t\t{entry['duration']}").bold = True
        # Add Role
        doc.add_paragraph(f"Role: {entry['role']}").bold = True
        # Split responsibilities by line breaks and add as bullet points
        responsibilities = entry['responsibilities'].split('\n')
        header4=doc.add_heading("Responsibilities:", level=3)
        for run in header4.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        # Create a new paragraph after the heading
        paragraph = doc.add_paragraph()
        # Add a blank line to the new paragraph
        paragraph.paragraph_format.space_after = Pt(6)
        for responsibility in responsibilities:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run("• ").bold = True
            p.add_run(responsibility)
            p.paragraph_format.line_spacing = Pt(9)
            
        # Update the flag
        first_entry = False

    # Add Education section
    header5=doc.add_heading("Education", level=2)
    for run in header5.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    # Create a new paragraph after the heading
    paragraph = doc.add_paragraph()
    # Add a blank line to the new paragraph
    paragraph.paragraph_format.space_after = Pt(9)
    education_paragraph = doc.add_paragraph()
    education_paragraph.paragraph_format.left_indent = Inches(0.3)
    education_paragraph.add_run("• ").bold = True
    education_paragraph.add_run(context["education"])
    education_paragraph.paragraph_format.line_spacing = Pt(9)



    # Specify the directory where you want to save the file
    output_directory = "templates/temp-docx"

    # Ensure the output directory exists, create it if necessary
    os.makedirs(output_directory, exist_ok=True)

    # Specify the full path for the output file
    output_file_path = os.path.join(output_directory, "generated_resume.docx")

    # Save the Word document
    doc.save(output_file_path)

    return FileResponse(output_file_path, headers={"Content-Disposition": "attachment; filename=generated_resume.docx"})
    
# Route to render the create resume form
@app.get("/create_resume_form/")
async def create_resume_form(request: Request):
    return templates.TemplateResponse("resume_form.html", {"request": request})
